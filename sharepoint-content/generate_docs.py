"""Generate all HR SharePoint policy documents as Word (.docx) files.

Run: python generate_docs.py
Creates documents in subfolders:
  HR-Policies/
  Benefits-Insurance/
  Payroll-Tax/
  Leave-TimeOff/
  Personal-Data-Changes/
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))


def _style_doc(doc: Document):
    """Apply consistent styling to the document."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


def _add_metadata_table(doc, rows: list[tuple[str, str]]):
    """Add a metadata table at the top of the document."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.style.font.size = Pt(10)
    doc.add_paragraph()


def _add_table(doc, headers: list[str], rows: list[list[str]]):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        for p in table.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    doc.add_paragraph()


def _save(doc: Document, folder: str, filename: str):
    path = os.path.join(BASE, folder)
    os.makedirs(path, exist_ok=True)
    filepath = os.path.join(path, filename)
    doc.save(filepath)
    print(f"  ✓ {folder}/{filename}")


# ═══════════════════════════════════════════════════════════════════════════════
# 1. HR POLICIES
# ═══════════════════════════════════════════════════════════════════════════════

def create_grievance_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("GRV-001: Formal Grievance Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "GRV-001"),
        ("Effective Date", "January 1, 2025"),
        ("Last Reviewed", "January 15, 2026"),
        ("Owner", "Chief Human Resources Officer"),
        ("Applies To", "All employees (full-time, part-time, and temporary)"),
        ("Classification", "Internal – Employee Confidential"),
    ])

    doc.add_heading("1. Purpose", level=2)
    doc.add_paragraph(
        "This policy establishes a formal process for employees to raise serious workplace concerns "
        "that cannot be resolved through informal channels. The grievance process ensures that employee "
        "complaints are handled fairly, consistently, and in a timely manner while protecting the rights "
        "of all parties involved."
    )

    doc.add_heading("2. Scope", level=2)
    doc.add_heading("2.1 What Constitutes a Formal Grievance", level=3)
    doc.add_paragraph(
        "A formal grievance is a written complaint regarding a serious workplace issue that falls "
        "into one or more of the following categories:"
    )
    items = [
        ("Discrimination", "Adverse treatment based on race, color, religion, sex (including pregnancy, sexual orientation, and gender identity), national origin, age (40+), disability, genetic information, veteran status, or any other characteristic protected by applicable law."),
        ("Harassment", "Unwelcome conduct based on a protected characteristic that creates a hostile, intimidating, or offensive work environment, or when submission to such conduct is made a condition of employment."),
        ("Sexual Harassment", "Unwelcome sexual advances, requests for sexual favors, or other verbal or physical conduct of a sexual nature that affects employment decisions or creates a hostile environment."),
        ("Retaliation", "Adverse action taken against an employee for reporting concerns, participating in an investigation, or exercising rights under company policy or law."),
        ("Bullying", "Repeated, unreasonable behavior directed toward an employee that creates a risk to health, safety, or well-being. This includes verbal abuse, intimidation, humiliation, or sabotage that interferes with work."),
        ("Hostile Work Environment", "Pervasive conduct based on protected characteristics that is severe enough to create an intimidating, hostile, or abusive working environment."),
        ("Significant Policy Violations", "Material breaches of company policy that affect employee rights, safety, or working conditions."),
        ("Wage and Hour Violations", "Disputes regarding compensation, overtime pay, or benefits entitlements."),
        ("Unsafe Working Conditions", "Concerns about workplace safety that have not been addressed through normal safety reporting channels."),
    ]
    for title, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)
        p.style = doc.styles["List Bullet"]

    doc.add_heading("2.2 What Does NOT Constitute a Formal Grievance", level=3)
    doc.add_paragraph(
        "The following types of concerns should be addressed through alternative channels (see Section 8):"
    )
    non_items = [
        "Minor interpersonal disagreements – One-time conflicts between colleagues without discriminatory or harassing elements.",
        "Management style preferences – Disagreements with a manager's approach that do not involve discrimination, harassment, or policy violations.",
        "Day-to-day workplace annoyances – Issues such as noise levels, temperature preferences, kitchen cleanliness, parking disputes, or personal space concerns.",
        "Performance feedback disagreements – Disputes about performance ratings should be addressed through the Performance Review Appeals process (PER-003).",
        "Compensation disputes – General pay dissatisfaction should be directed to the Total Rewards team unless the concern involves discriminatory pay practices.",
        "Scheduling preferences – Work schedule requests should be handled through the Flexible Work Arrangements policy (WRK-005).",
        "Facilities complaints – Issues with office equipment, workspace conditions, or building maintenance should be reported to Facilities Management.",
    ]
    for item in non_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Confidentiality", level=2)
    doc.add_paragraph("All grievances are handled under strict confidentiality:")
    conf_items = [
        "Information is shared only with individuals who have a legitimate need to know in order to investigate and resolve the matter.",
        "All investigation files are maintained separately from personnel files.",
        "Electronic records are stored in access-restricted systems with audit logging.",
        "Breach of grievance confidentiality by any party may result in disciplinary action.",
        "Employees may request anonymity, though this may limit the organization's ability to fully investigate.",
    ]
    for item in conf_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Grievance Process", level=2)

    doc.add_heading("Step 1: Intake and Classification (Days 1-2)", level=3)
    steps1 = [
        "Employee submits a grievance through any of the following channels: HR Concierge (online portal), direct submission to Employee Relations, email to grievance@company.com, or in-person meeting with any HR representative.",
        "Employee Relations performs initial classification within 24 hours of receipt.",
        "The grievance is assigned a case number (format: GRV-YYYYMMDD-XXXX) and a primary ER specialist.",
        "Employee receives written acknowledgment within 48 hours confirming receipt and next steps.",
    ]
    for s in steps1:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Step 2: Investigation (Days 3-15)", level=3)
    steps2 = [
        "The assigned ER specialist conducts a thorough, impartial investigation including: detailed interview with the complainant, interview with the respondent (accused party), interviews with relevant witnesses, review of documentary evidence (emails, messages, records), and review of applicable policies and precedents.",
        "All parties have the right to be accompanied by a colleague (not acting as legal counsel) during interviews.",
        "Investigation notes are documented in the case management system.",
        "The target completion for investigations is 10 business days, though complex cases may require extensions.",
    ]
    for s in steps2:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Step 3: Resolution and Outcome (Days 16-20)", level=3)
    doc.add_paragraph(
        "The ER specialist prepares a written findings report. The report is reviewed by the Employee Relations Manager "
        "and, for high-severity cases, by Legal. Appropriate remedial actions are determined, which may include:"
    )
    remedies = [
        "Mediation between parties",
        "Coaching or training for the respondent",
        "Formal warning or disciplinary action",
        "Policy or process changes",
        "Transfer or reassignment (with employee consent)",
        "Termination of the respondent (in severe cases)",
    ]
    for r in remedies:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_paragraph(
        "The complainant is informed of the outcome and any actions taken (to the extent permitted "
        "by confidentiality obligations to other parties)."
    )

    doc.add_heading("Step 4: Follow-Up (Days 21-30)", level=3)
    doc.add_paragraph(
        "ER specialist checks in with the complainant to ensure the resolution is effective. "
        "If the complainant is dissatisfied with the outcome, they may escalate to the Appeals process (see Section 6). "
        "The case file is closed and retained per the Records Retention policy (minimum 7 years)."
    )

    doc.add_heading("5. Non-Retaliation Protection", level=2)
    doc.add_paragraph("The company strictly prohibits retaliation against any employee who:")
    retal_items = [
        "Files a grievance in good faith",
        "Participates in a grievance investigation as a witness",
        "Provides information related to a grievance",
        "Opposes practices they reasonably believe to be unlawful",
    ]
    for r in retal_items:
        doc.add_paragraph(r, style="List Bullet")
    p = doc.add_paragraph()
    run = p.add_run(
        "Any employee found to have retaliated against a grievant or witness will be subject to "
        "disciplinary action up to and including termination."
    )
    run.bold = True

    doc.add_heading("6. Appeals Process", level=2)
    doc.add_paragraph(
        "Level 1 Appeal: Submit a written appeal to the Director of Employee Relations within 10 business "
        "days of receiving the outcome. The appeal will be reviewed by a different ER specialist."
    )
    doc.add_paragraph(
        "Level 2 Appeal: If still dissatisfied, submit a final appeal to the VP of Human Resources within "
        "10 business days. This decision is final."
    )

    doc.add_heading("7. Service Level Agreements", level=2)
    _add_table(doc,
        ["Milestone", "Target Timeline"],
        [
            ["Acknowledgment of receipt", "48 hours"],
            ["Initial classification", "24 hours"],
            ["Investigation commencement", "3 business days"],
            ["Investigation completion", "10 business days"],
            ["Findings report", "15 business days"],
            ["Resolution communication", "20 business days"],
            ["Follow-up check-in", "30 business days"],
        ]
    )

    doc.add_heading("8. Alternative Resolution Channels", level=2)
    doc.add_paragraph(
        "For concerns that do not meet the threshold for a formal grievance, the following resources are available:"
    )
    alt = [
        "Manager Mediation: Speak with your direct manager or skip-level manager about interpersonal issues.",
        "Peer Mediation Program: Trained peer mediators can facilitate conversations between colleagues.",
        "Employee Assistance Program (EAP): Confidential counseling and support services available 24/7 at 1-800-EAP-HELP.",
        "HR Business Partner: Your designated HRBP can provide guidance and informal support.",
        "Open Door Policy: Any employee may raise concerns with any level of management.",
        "Ethics Hotline: Anonymous reporting available at 1-800-ETHICS-1 or ethics.company.com.",
    ]
    for a in alt:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("9. Record Keeping", level=2)
    doc.add_paragraph(
        "All grievance records are retained for a minimum of 7 years from the date of case closure. "
        "Records include: intake form, investigation notes, interview records, evidence, findings report, "
        "and resolution documentation. Access to grievance records is restricted to authorized ER personnel and Legal."
    )

    doc.add_paragraph()
    doc.add_paragraph("For questions about this policy, contact Employee Relations at er@company.com or ext. 5500.").italic = True

    _save(doc, "HR-Policies", "GRV-001-Formal-Grievance-Policy.docx")


def create_code_of_conduct():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("COC-001: Code of Conduct", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "COC-001"),
        ("Effective Date", "January 1, 2025"),
        ("Last Reviewed", "March 1, 2026"),
        ("Owner", "Chief Human Resources Officer & General Counsel"),
        ("Applies To", "All employees, contractors, and temporary workers"),
        ("Classification", "Internal – All Employees"),
    ])

    doc.add_heading("1. Purpose", level=2)
    doc.add_paragraph(
        "This Code of Conduct establishes the standards of behavior expected of all employees. "
        "It serves as a guide for making ethical decisions and fostering a respectful, inclusive, "
        "and productive workplace."
    )

    doc.add_heading("2. Core Values", level=2)
    values = [
        ("Respect", "We treat every individual with dignity, regardless of their role, background, or identity."),
        ("Integrity", "We act honestly, transparently, and in compliance with all applicable laws and regulations."),
        ("Accountability", "We take responsibility for our actions and their impact on others."),
        ("Inclusion", "We value diverse perspectives and create an environment where everyone can thrive."),
    ]
    for title, desc in values:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("3. Professional Conduct Standards", level=2)

    doc.add_heading("3.1 Respectful Workplace", level=3)
    for item in [
        "Communicate professionally and courteously in all interactions",
        "Listen actively and consider others' perspectives",
        "Avoid gossip, rumors, or statements that could harm another's reputation",
        "Respect physical and personal boundaries",
        "Use inclusive language free from slurs, stereotypes, or derogatory terms",
        "Address disagreements constructively and respectfully",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 Conflicts of Interest", level=3)
    for item in [
        "Disclose any actual or potential conflicts of interest to your manager and HR",
        "Do not use your position for personal gain at the company's expense",
        "Obtain approval before engaging in outside employment that may conflict with your role",
        "Do not accept gifts or entertainment that could influence business decisions (threshold: $100)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Prohibited Conduct", level=2)

    doc.add_heading("4.1 Discrimination", level=3)
    doc.add_paragraph(
        "Any adverse treatment of an employee based on race, color, ethnicity, national origin, sex, gender, "
        "gender identity, gender expression, sexual orientation, age (40+), religion, disability, genetic "
        "information, veteran status, pregnancy, marital status, or any other protected characteristic."
    )

    doc.add_heading("4.2 Harassment", level=3)
    doc.add_paragraph("Unwelcome conduct based on a protected characteristic, including:")
    for item in [
        "Verbal harassment: slurs, epithets, offensive jokes, ridicule, mockery",
        "Physical harassment: unwanted touching, physical intimidation, blocking movement",
        "Visual harassment: offensive images, posters, emails, or social media posts",
        "Sexual harassment: unwanted sexual advances, requests for sexual favors",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4.3 Bullying and Intimidation", level=3)
    for item in [
        "Repeated aggressive behavior intended to intimidate, degrade, or humiliate",
        "Shouting, cursing, or using threatening language",
        "Deliberate exclusion from meetings, communications, or social activities with intent to ostracize",
        "Sabotaging another's work or reputation",
        "Spreading malicious rumors",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Reporting Concerns", level=2)
    doc.add_paragraph("Employees who witness or experience violations should report through:")
    _add_table(doc,
        ["Channel", "Contact"],
        [
            ["Direct Manager", "For minor concerns or first-time issues"],
            ["HR Business Partner", "For guidance on appropriate next steps"],
            ["Employee Relations", "er@company.com / ext. 5500"],
            ["Ethics Hotline", "1-800-ETHICS-1 / ethics.company.com (anonymous)"],
            ["HR Concierge", "Online portal for guided submission"],
        ]
    )

    doc.add_heading("6. Disciplinary Framework", level=2)
    _add_table(doc,
        ["Step", "Action"],
        [
            ["1", "Verbal Counseling – For minor first-time infractions"],
            ["2", "Written Warning – For repeated minor infractions or moderate violations"],
            ["3", "Final Written Warning – For serious violations or continued misconduct"],
            ["4", "Suspension – Pending investigation of serious allegations"],
            ["5", "Termination – For severe violations or gross misconduct"],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph("For questions, contact HR at hr@company.com or your HR Business Partner.").italic = True

    _save(doc, "HR-Policies", "COC-001-Code-of-Conduct.docx")


def create_anti_discrimination_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("EEO-001: Anti-Discrimination & Equal Employment Opportunity Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "EEO-001"),
        ("Effective Date", "January 1, 2025"),
        ("Last Reviewed", "February 1, 2026"),
        ("Owner", "Chief Human Resources Officer & General Counsel"),
        ("Applies To", "All employees, applicants, contractors, and vendors"),
        ("Classification", "Internal – All Employees"),
    ])

    doc.add_heading("1. Policy Statement", level=2)
    doc.add_paragraph(
        "The company is committed to providing equal employment opportunities to all employees and applicants "
        "without regard to race, color, religion, sex (including pregnancy, sexual orientation, and gender identity), "
        "national origin, age (40 and over), disability, genetic information, veteran status, or any other "
        "characteristic protected by applicable federal, state, or local law."
    )

    doc.add_heading("2. Protected Characteristics", level=2)
    _add_table(doc,
        ["Protected Characteristic", "Applicable Law"],
        [
            ["Race, color, ethnicity", "Title VII of the Civil Rights Act"],
            ["Sex, gender, gender identity", "Title VII, as amended"],
            ["Sexual orientation", "Title VII (Bostock v. Clayton County)"],
            ["National origin", "Title VII"],
            ["Religion or creed", "Title VII"],
            ["Age (40 and over)", "Age Discrimination in Employment Act (ADEA)"],
            ["Disability (physical or mental)", "Americans with Disabilities Act (ADA)"],
            ["Genetic information", "Genetic Information Nondiscrimination Act (GINA)"],
            ["Pregnancy, childbirth", "Pregnancy Discrimination Act (PDA)"],
            ["Veteran / military status", "USERRA, VEVRAA"],
        ]
    )

    doc.add_heading("3. Reasonable Accommodations", level=2)
    doc.add_heading("3.1 Disability Accommodations", level=3)
    doc.add_paragraph(
        "The company will provide reasonable accommodations to qualified individuals with disabilities, "
        "unless doing so would create an undue hardship. Accommodations may include:"
    )
    for item in [
        "Modified work schedules or duties",
        "Assistive technology or equipment",
        "Accessible facilities",
        "Leave for treatment",
        "Reassignment to vacant positions",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 Religious Accommodations", level=3)
    doc.add_paragraph(
        "The company will reasonably accommodate employees' sincerely held religious beliefs, "
        "including schedule modifications, dress code exceptions, and space for prayer or meditation."
    )

    doc.add_heading("3.3 Pregnancy Accommodations", level=3)
    doc.add_paragraph(
        "In accordance with the Pregnant Workers Fairness Act, the company provides reasonable accommodations "
        "for known limitations related to pregnancy, childbirth, or related medical conditions."
    )

    doc.add_heading("4. Discrimination Indicators", level=2)
    indicators = [
        ("Disparate treatment", "Treating employees differently based on a protected characteristic"),
        ("Disparate impact", "Neutral policies that disproportionately affect a protected group without business justification"),
        ("Stereotyping", "Making assumptions about abilities or behavior based on group membership"),
        ("Exclusion", "Systematically leaving individuals out of meetings, projects, or events based on protected characteristics"),
        ("Retaliation", "Punishing employees who report discrimination or participate in investigations"),
    ]
    for title, desc in indicators:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("5. Age Discrimination", level=2)
    doc.add_paragraph("The company prohibits the following age-related discriminatory practices:")
    for item in [
        "Using age as a factor in hiring, promotion, or termination decisions",
        "Mandatory retirement based on age (except as permitted by law)",
        "Age-related harassment including jokes, comments, or ridicule",
        "Job advertisements that specify age preferences (e.g., 'young and energetic')",
        "Pressuring older employees to retire or accept early retirement",
        "Excluding older employees from training or advancement opportunities",
        "Using age as a factor in reduction-in-force decisions",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. Reporting Discrimination", level=2)
    _add_table(doc,
        ["Channel", "Contact"],
        [
            ["HR Business Partner", "For guidance and informal resolution"],
            ["Employee Relations", "er@company.com / ext. 5500"],
            ["Ethics Hotline", "1-800-ETHICS-1 (anonymous)"],
            ["HR Concierge", "Online portal"],
            ["EEOC", "External — eeoc.gov"],
        ]
    )

    doc.add_heading("7. Manager Responsibilities", level=2)
    for item in [
        "Model inclusive behavior and language",
        "Make employment decisions based solely on legitimate business factors",
        "Address discriminatory behavior immediately",
        "Participate in required anti-discrimination training",
        "Document employment decisions with clear, objective justification",
        "Consult HR before making adverse employment decisions",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("For questions, contact Employee Relations at er@company.com or the Legal department.").italic = True

    _save(doc, "HR-Policies", "EEO-001-Anti-Discrimination-Policy.docx")


def create_harassment_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("HAR-001: Workplace Harassment Prevention Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "HAR-001"),
        ("Effective Date", "January 1, 2025"),
        ("Last Reviewed", "February 15, 2026"),
        ("Owner", "Chief Human Resources Officer"),
        ("Applies To", "All employees, contractors, interns, and visitors"),
        ("Classification", "Internal – All Employees"),
    ])

    doc.add_heading("1. Policy Statement", level=2)
    doc.add_paragraph(
        "The company is committed to maintaining a workplace free from harassment of any kind. "
        "Harassment based on any protected characteristic is illegal, violates company policy, and will not be tolerated. "
        "This policy applies to all work-related settings including the workplace, company events, business travel, "
        "and electronic communications."
    )

    doc.add_heading("2. Definition of Harassment", level=2)
    doc.add_heading("2.1 General Harassment", level=3)
    doc.add_paragraph(
        "Harassment is unwelcome conduct based on a protected characteristic that: (a) is made a condition of "
        "employment (quid pro quo), OR (b) is sufficiently severe or pervasive to create a hostile, intimidating, "
        "or offensive work environment."
    )

    doc.add_heading("2.2 Sexual Harassment", level=3)
    doc.add_paragraph("Examples of sexual harassment include but are not limited to:")
    for item in [
        "Unwanted sexual comments, jokes, or innuendo",
        "Displaying sexually explicit materials",
        "Unwanted physical contact (touching, hugging, kissing)",
        "Sexually suggestive gestures or staring",
        "Sending sexual messages, images, or links",
        "Pressuring for dates or sexual activity",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.3 Bullying", level=3)
    doc.add_paragraph("Workplace bullying includes:")
    for item in [
        "Repeated verbal abuse (yelling, cursing, belittling)",
        "Humiliation in front of colleagues",
        "Deliberate exclusion or isolation",
        "Undermining or sabotaging work",
        "Spreading malicious rumors",
        "Excessive, unreasonable criticism or monitoring",
        "Taking credit for another's work",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. What Harassment Is NOT", level=2)
    for item in [
        "Legitimate performance management and feedback, even if uncomfortable",
        "Reasonable workplace directives and expectations",
        "Professional disagreements or debates conducted respectfully",
        "Isolated minor incidents (though these may still warrant coaching)",
        "Management decisions made in good faith",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Reporting Procedures", level=2)
    _add_table(doc,
        ["Channel", "Contact", "Notes"],
        [
            ["Direct Manager", "Your manager", "Unless the manager is the harasser"],
            ["HR Business Partner", "Your HRBP", "Confidential guidance"],
            ["Employee Relations", "er@company.com / ext. 5500", "Formal investigation"],
            ["Ethics Hotline", "1-800-ETHICS-1", "Anonymous option"],
            ["HR Concierge", "Online portal", "Guided submission"],
        ]
    )

    doc.add_heading("5. Manager Obligations", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Managers who witness harassment or receive a complaint must take immediate action:")
    run.bold = True
    for item in [
        "Ensure the employee feels safe and remove them from the immediate situation if necessary",
        "Report the incident to Employee Relations within 24 hours — even if the employee asks them not to",
        "Do not attempt to investigate or resolve the matter independently",
        "Maintain confidentiality",
        "Monitor for retaliation against the reporting employee",
    ]:
        doc.add_paragraph(item, style="List Number")

    p = doc.add_paragraph()
    run = p.add_run("Failure to report known harassment may result in disciplinary action against the manager.")
    run.bold = True

    doc.add_heading("6. Investigation and Remedial Actions", level=2)
    doc.add_paragraph("All complaints are investigated per the Formal Grievance Policy (GRV-001). Remedial actions may include:")
    for item in [
        "Mandatory harassment prevention training",
        "Formal warning or suspension",
        "Demotion or transfer",
        "Termination of employment",
        "Referral to law enforcement (criminal conduct)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("For questions, contact Employee Relations at er@company.com or ext. 5500.").italic = True

    _save(doc, "HR-Policies", "HAR-001-Workplace-Harassment-Policy.docx")


def create_retaliation_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("RET-001: Retaliation Protection Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "RET-001"),
        ("Effective Date", "January 1, 2025"),
        ("Last Reviewed", "January 15, 2026"),
        ("Owner", "Chief Human Resources Officer & General Counsel"),
        ("Applies To", "All employees, managers, and supervisors"),
        ("Classification", "Internal – All Employees"),
    ])

    doc.add_heading("1. Policy Statement", level=2)
    doc.add_paragraph(
        "The company strictly prohibits retaliation against any employee who, in good faith, reports a violation "
        "of company policy, participates in an investigation, files a complaint with an external agency, or exercises "
        "any right protected by law."
    )

    doc.add_heading("2. Protected Activities", level=2)
    for item in [
        "Filing a grievance or complaint under any company policy",
        "Reporting discrimination, harassment, fraud, safety violations, or ethical concerns",
        "Participating as a witness in an internal investigation",
        "Filing a charge with a government agency (EEOC, OSHA, DOL, SEC, etc.)",
        "Refusing to participate in activities they reasonably believe are unlawful",
        "Requesting a reasonable accommodation for a disability or religious practice",
        "Exercising rights under FMLA, ADA, FLSA, or other employment laws",
        "Engaging in whistleblower activities",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Definition of Retaliation", level=2)
    _add_table(doc,
        ["Category", "Examples"],
        [
            ["Employment status", "Termination, demotion, suspension, forced transfer"],
            ["Compensation", "Pay reduction, denial of bonus, reduced hours"],
            ["Work conditions", "Undesirable schedule changes, reassignment to menial tasks"],
            ["Career impact", "Denial of promotion, negative review, removal from training"],
            ["Interpersonal", "Harassment, intimidation, threats, hostile behavior"],
            ["Social exclusion", "Deliberate isolation, excluding from meetings or team events"],
        ]
    )

    doc.add_heading("4. Reporting Retaliation", level=2)
    doc.add_paragraph(
        "Employees who believe they are experiencing retaliation should document the adverse actions "
        "and report to Employee Relations at er@company.com or ext. 5500. The Ethics Hotline (1-800-ETHICS-1) "
        "is also available for anonymous reporting."
    )

    doc.add_heading("5. Consequences", level=2)
    doc.add_paragraph(
        "If retaliation is substantiated, the retaliator will face disciplinary action up to and including termination. "
        "The company will take steps to restore the affected employee to their prior position, compensation, and standing."
    )

    doc.add_paragraph()
    doc.add_paragraph("For questions, contact Employee Relations at er@company.com or the Legal department.").italic = True

    _save(doc, "HR-Policies", "RET-001-Retaliation-Protection-Policy.docx")


def create_employee_relations_procedures():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("ER-001: Employee Relations Procedures", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "ER-001"),
        ("Effective Date", "January 1, 2025"),
        ("Last Reviewed", "March 1, 2026"),
        ("Owner", "Director of Employee Relations"),
        ("Applies To", "HR team members, managers, and all employees"),
        ("Classification", "Internal – All Employees"),
    ])

    doc.add_heading("1. ER Services", level=2)
    _add_table(doc,
        ["Service", "Description", "Timeline"],
        [
            ["Coaching", "One-on-one guidance for workplace challenges", "As needed"],
            ["Mediation", "Facilitated conversation between parties", "Within 5 business days"],
            ["Team Intervention", "Session for teams experiencing dysfunction", "Within 10 business days"],
            ["Skip-Level Meeting", "Discussion with a higher-level manager", "Within 3 business days"],
        ]
    )

    doc.add_heading("2. When to Contact ER vs. Other Resources", level=2)
    _add_table(doc,
        ["Situation", "Resource"],
        [
            ["Discrimination, harassment, or retaliation", "Employee Relations"],
            ["Interpersonal conflict with a colleague", "Manager first, then ER if unresolved"],
            ["Performance concerns about a direct report", "HR Business Partner + ER"],
            ["Personal stress, mental health, substance abuse", "Employee Assistance Program (EAP)"],
            ["Pay or benefits questions", "Total Rewards / Payroll"],
            ["Safety concern", "Safety Department / OSHA"],
            ["Policy question", "HR Business Partner or HR Concierge"],
        ]
    )

    doc.add_heading("3. Case Categories and SLAs", level=2)
    _add_table(doc,
        ["Category", "Description", "SLA"],
        [
            ["Critical", "Imminent safety threat, sexual assault, threats of violence", "Immediate response, same day"],
            ["High", "Discrimination, harassment, retaliation", "Investigation begins within 3 business days"],
            ["Medium", "Policy violations, conduct issues, workplace conflict", "Response within 5 business days"],
            ["Low", "General guidance, policy clarification, coaching", "Response within 7 business days"],
        ]
    )

    doc.add_heading("4. Employee Assistance Program (EAP)", level=2)
    doc.add_paragraph("The EAP provides confidential, no-cost support for employees and their household members:")
    for item in [
        "Mental health counseling (up to 8 sessions per issue per year)",
        "Substance abuse assessment and referral",
        "Financial counseling and planning",
        "Legal consultation (30-minute initial consultation)",
        "Work-life resources (childcare, eldercare, moving assistance)",
        "Critical incident response",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("Contact: 1-800-EAP-HELP (available 24/7/365) | Website: eap.company.com")

    doc.add_heading("5. Contact Information", level=2)
    _add_table(doc,
        ["Team", "Contact"],
        [
            ["Employee Relations", "er@company.com / ext. 5500"],
            ["Ethics Hotline", "1-800-ETHICS-1 / ethics.company.com"],
            ["EAP", "1-800-EAP-HELP / eap.company.com"],
            ["Security", "ext. 9911 (emergency)"],
            ["HR Business Partners", "hrbp@company.com"],
        ]
    )

    _save(doc, "HR-Policies", "ER-001-Employee-Relations-Procedures.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. BENEFITS & INSURANCE
# ═══════════════════════════════════════════════════════════════════════════════

def create_health_insurance_guide():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("BEN-001: Health Insurance Plan Guide", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "BEN-001"),
        ("Plan Year", "January 1 – December 31, 2026"),
        ("Owner", "VP of Total Rewards"),
        ("Applies To", "Benefits-eligible employees (30+ hours/week)"),
    ])

    doc.add_heading("1. Eligibility", level=2)
    doc.add_paragraph("Full-time employees (30+ hours/week). Coverage begins 1st of the month following hire date.")
    doc.add_heading("Eligible Dependents", level=3)
    for item in [
        "Legal spouse (marriage certificate required)",
        "Domestic partner (affidavit required)",
        "Dependent children up to age 26 (biological, adopted, step, foster, legal guardianship)",
        "Disabled adult dependents (age 26+) who became disabled before age 26",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Plan Options", level=2)

    doc.add_heading("2.1 PPO Plan (Preferred Provider Organization)", level=3)
    _add_table(doc,
        ["Feature", "In-Network", "Out-of-Network"],
        [
            ["Monthly Premium (EE only)", "$185", "—"],
            ["Monthly Premium (EE + Spouse)", "$420", "—"],
            ["Monthly Premium (EE + Children)", "$370", "—"],
            ["Monthly Premium (Family)", "$595", "—"],
            ["Annual Deductible (Individual)", "$500", "$1,500"],
            ["Annual Deductible (Family)", "$1,000", "$3,000"],
            ["Out-of-Pocket Max (Individual)", "$4,000", "$8,000"],
            ["Primary Care Visit", "$25 copay", "40% after deductible"],
            ["Specialist Visit", "$40 copay", "40% after deductible"],
            ["Preventive Care", "$0 (100% covered)", "40% after deductible"],
            ["Emergency Room", "$200 copay", "$200 copay"],
            ["Rx – Generic", "$10 copay", "$25 copay"],
            ["Rx – Brand Preferred", "$35 copay", "$60 copay"],
            ["Rx – Specialty", "20% up to $200", "30% up to $300"],
            ["Mental Health – Outpatient", "$25 copay", "40% after deductible"],
        ]
    )
    doc.add_paragraph("Best for: Employees who want flexibility to see any provider without referrals.")

    doc.add_heading("2.2 HDHP Plan with HSA", level=3)
    _add_table(doc,
        ["Feature", "In-Network", "Out-of-Network"],
        [
            ["Monthly Premium (EE only)", "$95", "—"],
            ["Monthly Premium (Family)", "$320", "—"],
            ["Annual Deductible (Individual)", "$1,600", "$3,200"],
            ["Annual Deductible (Family)", "$3,200", "$6,400"],
            ["Out-of-Pocket Max (Individual)", "$5,000", "$10,000"],
            ["All Services After Deductible", "10% coinsurance", "30% coinsurance"],
            ["Preventive Care", "$0 (100% covered)", "30% after deductible"],
            ["HSA Employer Contribution", "$750 (EE) / $1,500 (Family)", "—"],
        ]
    )
    doc.add_paragraph("HSA Contribution Limits (2026): Individual $4,300 | Family $8,550 | Catch-up (55+) additional $1,000")
    doc.add_paragraph("Best for: Healthy employees who want lower premiums and tax-free HSA savings.")

    doc.add_heading("2.3 HMO Plan", level=3)
    _add_table(doc,
        ["Feature", "In-Network Only"],
        [
            ["Monthly Premium (EE only)", "$145"],
            ["Monthly Premium (Family)", "$475"],
            ["Annual Deductible (Individual)", "$250"],
            ["Out-of-Pocket Max (Individual)", "$3,500"],
            ["Primary Care Visit", "$20 copay"],
            ["Specialist Visit (requires referral)", "$35 copay"],
            ["Emergency Room", "$150 copay"],
            ["Rx – Generic", "$5 copay"],
        ]
    )
    doc.add_paragraph("Note: HMO requires a PCP and referrals for specialists. Out-of-network not covered except emergencies.")

    doc.add_heading("3. Qualifying Life Events", level=2)
    doc.add_paragraph("You may make changes within 30 days of a qualifying life event:")
    for item in [
        "Marriage or divorce",
        "Birth, adoption, or placement of a child",
        "Death of a spouse or dependent",
        "Loss of other coverage (spouse's plan, COBRA exhaustion, aging out)",
        "Gain of other coverage",
        "Change in employment status affecting eligibility",
        "Court order requiring coverage",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    p = doc.add_paragraph()
    run = p.add_run("Important: You must notify HR and submit changes within 30 days of the event. Supporting documentation is required.")
    run.bold = True

    doc.add_heading("4. COBRA Continuation", level=2)
    doc.add_paragraph(
        "If you lose coverage due to termination or reduction in hours, you and your dependents may continue "
        "coverage under COBRA for up to 18 months (36 months for certain events) at full cost plus 2% admin fee."
    )

    _save(doc, "Benefits-Insurance", "BEN-001-Health-Insurance-Plan-Guide.docx")


def create_dental_vision_guide():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("BEN-002: Dental & Vision Benefits Guide", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "BEN-002"),
        ("Plan Year", "January 1 – December 31, 2026"),
        ("Owner", "VP of Total Rewards"),
    ])

    doc.add_heading("1. Dental Plan", level=2)
    _add_table(doc,
        ["Feature", "In-Network", "Out-of-Network"],
        [
            ["Monthly Premium (EE)", "$28", "—"],
            ["Monthly Premium (Family)", "$85", "—"],
            ["Annual Deductible", "$50 individual", "$50 individual"],
            ["Annual Maximum", "$2,000", "$1,500"],
            ["Preventive (cleanings, X-rays)", "100% covered", "80%"],
            ["Basic (fillings, extractions)", "80% after deductible", "60%"],
            ["Major (crowns, bridges, root canals)", "50% after deductible", "40%"],
            ["Orthodontia (children under 19)", "50% up to $2,000 lifetime max", "Not covered"],
            ["Orthodontia (adults)", "50% up to $1,500 lifetime max", "Not covered"],
        ]
    )

    doc.add_heading("2. Vision Plan", level=2)
    _add_table(doc,
        ["Feature", "In-Network"],
        [
            ["Monthly Premium (EE)", "$12"],
            ["Monthly Premium (Family)", "$35"],
            ["Eye Exam", "$10 copay (once per year)"],
            ["Lenses", "$25 copay (once per year)"],
            ["Frames", "$150 allowance (once every 2 years)"],
            ["Contact Lenses", "$150 allowance in lieu of frames"],
            ["Laser Vision Correction", "15% discount at participating providers"],
        ]
    )

    _save(doc, "Benefits-Insurance", "BEN-002-Dental-Vision-Guide.docx")


def create_life_insurance_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("BEN-003: Life Insurance & AD&D Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "BEN-003"),
        ("Owner", "VP of Total Rewards"),
        ("Applies To", "All benefits-eligible employees"),
    ])

    doc.add_heading("1. Basic Life Insurance (Company-Paid)", level=2)
    doc.add_paragraph("The company provides basic life insurance at no cost to you:")
    for item in [
        "Coverage amount: 1x annual base salary (up to $500,000)",
        "Accidental Death & Dismemberment (AD&D): Additional 1x annual base salary",
        "Coverage begins on your first day of benefits eligibility",
        "No medical exam required",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Supplemental Life Insurance (Voluntary)", level=2)
    _add_table(doc,
        ["Coverage", "Details"],
        [
            ["Employee supplemental", "1x to 5x annual salary (up to $1,000,000)"],
            ["Spouse/partner", "$10,000 to $250,000 in $10,000 increments"],
            ["Child(ren)", "$5,000 or $10,000 per child"],
            ["Guaranteed issue (no medical exam)", "Up to 3x salary at initial enrollment"],
            ["Evidence of Insurability", "Required for amounts above guaranteed issue"],
        ]
    )

    doc.add_heading("3. Beneficiary Designation", level=2)
    doc.add_paragraph(
        "You must designate a beneficiary for your life insurance benefits. If no beneficiary is designated, "
        "proceeds will be paid according to the plan's default order: spouse, children, parents, estate."
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "Important: Update your beneficiary after any qualifying life event (marriage, divorce, birth of a child, death of a beneficiary). "
        "Beneficiary changes can be made at any time through the HR portal or by contacting Benefits."
    )
    run.bold = True

    _save(doc, "Benefits-Insurance", "BEN-003-Life-Insurance-Policy.docx")


def create_401k_guide():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("BEN-004: 401(k) Retirement Plan Summary", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "BEN-004"),
        ("Plan Year", "January 1 – December 31, 2026"),
        ("Owner", "VP of Total Rewards"),
    ])

    doc.add_heading("1. Eligibility", level=2)
    doc.add_paragraph("All employees age 21+ are eligible to participate after 90 days of employment. Auto-enrollment at 6% of pay.")

    doc.add_heading("2. Contribution Limits (2026)", level=2)
    _add_table(doc,
        ["Type", "Limit"],
        [
            ["Employee pre-tax/Roth", "$23,500"],
            ["Catch-up (age 50-59, 64+)", "Additional $7,500"],
            ["Super catch-up (age 60-63)", "Additional $11,250"],
            ["Combined employee + employer", "$70,000"],
        ]
    )

    doc.add_heading("3. Employer Match", level=2)
    doc.add_paragraph("The company matches 100% of the first 4% of pay and 50% of the next 2% of pay (up to 5% total match). "
                      "Vesting schedule: 100% vested after 3 years of service (cliff vesting).")

    doc.add_heading("4. Investment Options", level=2)
    for item in [
        "Target-date funds (default investment)",
        "Index funds (S&P 500, Total Bond, International)",
        "Actively managed funds",
        "Company stock (limited to 10% of balance)",
        "Self-directed brokerage window",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _save(doc, "Benefits-Insurance", "BEN-004-401k-Plan-Summary.docx")


def create_qualifying_life_events_guide():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("BEN-005: Qualifying Life Events Guide", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "BEN-005"),
        ("Owner", "VP of Total Rewards"),
        ("Applies To", "All benefits-eligible employees"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "A qualifying life event (QLE) allows you to make changes to your benefits elections outside of the annual "
        "Open Enrollment period. You must report the event and make changes within 30 days."
    )

    doc.add_heading("2. Qualifying Life Events", level=2)
    _add_table(doc,
        ["Event", "Allowed Changes", "Documentation Required", "Deadline"],
        [
            ["Marriage", "Add spouse; change plan tier; update beneficiaries", "Marriage certificate", "30 days"],
            ["Divorce / Legal separation", "Remove spouse; change plan tier; update beneficiaries", "Divorce decree / court order", "30 days"],
            ["Birth of a child", "Add child; change plan tier; increase life insurance", "Birth certificate", "30 days"],
            ["Adoption / Placement", "Add child; change plan tier", "Adoption papers / placement agreement", "30 days"],
            ["Death of a dependent", "Remove dependent; change plan tier; update beneficiaries", "Death certificate", "30 days"],
            ["Loss of other coverage", "Enroll or add dependents; change plans", "Letter from previous insurer with termination date", "30 days"],
            ["Gain of other coverage", "Drop coverage or dependents", "Proof of new coverage", "30 days"],
            ["Change in employment status", "Enroll, change, or drop coverage as applicable", "Employment documentation", "30 days"],
            ["Court order (QMCSO)", "Add child(ren) per court order", "Qualified Medical Child Support Order", "As ordered"],
        ]
    )

    doc.add_heading("3. How to Report a Qualifying Life Event", level=2)
    for i, step in enumerate([
        "Log in to the HR Portal or contact Benefits at benefits@company.com",
        "Select 'Report Life Event' and choose the event type",
        "Upload required documentation (see table above)",
        "Make your benefits elections — add/remove dependents, change plans, update beneficiaries",
        "Review and submit your changes",
        "You will receive a confirmation email within 2 business days",
    ], 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("4. Important Notes", level=2)
    for item in [
        "The 30-day deadline is strict. If you miss the window, you must wait until Open Enrollment.",
        "New coverage is effective on the date of the event (birth, marriage) or the 1st of the following month.",
        "Documentation must be submitted with your enrollment changes. Changes are not processed without documentation.",
        "If your life event affects tax withholding (e.g., marriage, birth), update your W-4 separately.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _save(doc, "Benefits-Insurance", "BEN-005-Qualifying-Life-Events-Guide.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 3. PAYROLL & TAX
# ═══════════════════════════════════════════════════════════════════════════════

def create_direct_deposit_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("PAY-001: Direct Deposit Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "PAY-001"),
        ("Owner", "Payroll Manager"),
        ("Applies To", "All employees receiving compensation"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "Direct deposit is the default payment method for all employee compensation. Employees may split deposits "
        "across up to 3 bank accounts."
    )

    doc.add_heading("2. Setting Up Direct Deposit", level=2)
    doc.add_paragraph("To set up or change your direct deposit:")
    for item in [
        "Log in to the HR Portal and navigate to Payroll > Direct Deposit",
        "Enter your bank name, routing number, and account number",
        "Select account type (Checking or Savings)",
        "Choose deposit allocation (Full, Partial Amount, or Remainder)",
        "Submit a void check or bank verification letter for new accounts",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("3. Security Requirements", level=2)
    doc.add_paragraph(
        "Bank account changes are classified as HIGH RISK due to payroll fraud concerns:"
    )
    for item in [
        "Identity verification is required for all bank changes",
        "A fraud prevention review may be triggered for certain patterns",
        "Changes take effect on the next pay cycle (not retroactive)",
        "You will receive email and SMS confirmation when changes are processed",
        "Two-factor authentication is required to access direct deposit settings",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Processing Timeline", level=2)
    _add_table(doc,
        ["Action", "Timeline"],
        [
            ["New direct deposit setup", "Effective next pay cycle"],
            ["Bank account change", "Effective next pay cycle (after verification)"],
            ["Adding a split deposit", "Effective next pay cycle"],
            ["Canceling direct deposit", "Paper check issued next pay cycle"],
        ]
    )

    _save(doc, "Payroll-Tax", "PAY-001-Direct-Deposit-Policy.docx")


def create_tax_withholding_guide():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("PAY-002: Tax Withholding Guide (W-4)", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "PAY-002"),
        ("Owner", "Payroll Manager"),
        ("Applies To", "All U.S. employees"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "Your W-4 form determines how much federal income tax is withheld from your paycheck. "
        "You should review and update your W-4 whenever you experience a life change that affects your tax situation."
    )

    doc.add_heading("2. When to Update Your W-4", level=2)
    for item in [
        "Getting married or divorced",
        "Having or adopting a child",
        "Starting a second job or your spouse starts working",
        "Receiving significant non-wage income",
        "Filing status changes",
        "Moving to a different state",
        "Annually, as part of tax planning",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Filing Status Options", level=2)
    _add_table(doc,
        ["Status", "Description"],
        [
            ["Single", "Unmarried, divorced, or legally separated"],
            ["Married Filing Jointly (MFJ)", "Married and filing a joint return"],
            ["Married Filing Separately (MFS)", "Married but filing separate returns"],
            ["Head of Household", "Unmarried and paying >50% of household expenses for a dependent"],
        ]
    )

    doc.add_heading("4. How to Update", level=2)
    doc.add_paragraph(
        "Update your W-4 through the HR Portal under Payroll > Tax Withholding. Changes take effect "
        "on the next pay cycle. You may also use the IRS Tax Withholding Estimator at irs.gov/W4App."
    )

    doc.add_heading("5. State Tax Implications for Relocations", level=2)
    doc.add_paragraph(
        "If you move to a different state, your state tax withholding may need to change. Some key considerations:"
    )
    for item in [
        "States with no income tax: AK, FL, NV, NH, SD, TN, TX, WA, WY",
        "Reciprocity agreements may apply if you live in one state and work in another",
        "Remote workers may owe taxes in both their home state and the state where the employer is located",
        "Notify Payroll of any address change so state withholding can be adjusted",
        "Multi-state tax situations may require Payroll to withhold for multiple states",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _save(doc, "Payroll-Tax", "PAY-002-Tax-Withholding-Guide.docx")


def create_payroll_change_sla():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("PAY-003: Payroll Change Processing SLA", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "PAY-003"),
        ("Owner", "Payroll Manager"),
    ])

    doc.add_heading("1. Processing Timelines", level=2)
    _add_table(doc,
        ["Change Type", "Risk Level", "Processing Time", "Approval Required"],
        [
            ["Address change (same state)", "Low", "2-3 business days", "No (self-service)"],
            ["Address change (cross-state)", "Medium", "5-7 business days", "Payroll review"],
            ["Direct deposit setup/change", "High", "Next pay cycle", "Identity verification"],
            ["W-4 / tax withholding change", "Medium", "Next pay cycle", "No (self-service)"],
            ["Legal name change", "High", "7-10 business days", "HR Ops + documentation"],
            ["Emergency contact update", "Low", "Immediate", "No (self-service)"],
            ["Preferred name change", "Low", "24-48 hours", "No (self-service)"],
            ["Beneficiary update", "Medium", "3-5 business days", "No"],
            ["New dependent addition", "Medium", "5-7 business days", "Documentation required"],
        ]
    )

    doc.add_heading("2. Payroll Calendar", level=2)
    doc.add_paragraph(
        "Pay dates are bi-weekly (every other Friday). Changes must be submitted by end of day Tuesday "
        "to be reflected in the following Friday's paycheck. Changes submitted after the cutoff will be "
        "effective in the next pay period."
    )

    _save(doc, "Payroll-Tax", "PAY-003-Payroll-Change-Processing-SLA.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 4. LEAVE & TIME OFF
# ═══════════════════════════════════════════════════════════════════════════════

def create_parental_leave_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("LVE-001: Parental Leave Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "LVE-001"),
        ("Effective Date", "January 1, 2025"),
        ("Last Reviewed", "January 1, 2026"),
        ("Owner", "VP of Total Rewards"),
        ("Applies To", "All benefits-eligible employees"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "The company provides generous parental leave to support employees welcoming a new child through birth, "
        "adoption, or foster placement. This policy applies to all eligible employees regardless of gender."
    )

    doc.add_heading("2. Leave Entitlements", level=2)
    _add_table(doc,
        ["Leave Type", "Duration", "Pay", "Eligibility"],
        [
            ["Primary caregiver (birth)", "16 weeks", "100% base pay", "All benefits-eligible employees"],
            ["Secondary caregiver", "8 weeks", "100% base pay", "All benefits-eligible employees"],
            ["Adoption / Foster placement", "12 weeks", "100% base pay", "All benefits-eligible employees"],
            ["Pregnancy-related medical leave", "Up to 4 additional weeks", "100% base pay", "Birth parent, with medical documentation"],
            ["FMLA (concurrent)", "Up to 12 weeks", "Per FMLA provisions", "Employees with 12+ months and 1,250+ hours"],
        ]
    )

    doc.add_heading("3. How to Apply", level=2)
    for i, step in enumerate([
        "Notify your manager and HR at least 30 days before the expected start date (or as soon as practical for unplanned events).",
        "Submit a leave request through the HR Portal under Leave > Parental Leave.",
        "Provide supporting documentation: expected due date (medical certification), adoption papers, or foster placement documentation.",
        "Meet with your manager to create a coverage plan for your responsibilities.",
        "HR will confirm your leave dates, pay continuation, and benefits continuation.",
    ], 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("4. Benefits During Leave", level=2)
    for item in [
        "Health insurance, dental, and vision continue at the same employee contribution rates",
        "Life insurance and disability coverage continue",
        "401(k) contributions are suspended during unpaid portions of leave (if any)",
        "PTO does not accrue during unpaid leave",
        "Company holidays falling during leave are paid separately",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Return to Work", level=2)
    doc.add_paragraph(
        "You are guaranteed reinstatement to the same or an equivalent position upon return from parental leave. "
        "If you need a modified schedule or additional time, discuss options with your manager and HR before your return date."
    )

    doc.add_heading("6. New Baby / Dependent Checklist", level=2)
    doc.add_paragraph("After the birth or placement of a new child, remember to:")
    for item in [
        "Add the child to your health insurance within 30 days (qualifying life event)",
        "Update your beneficiaries for life insurance and 401(k)",
        "Update your W-4 tax withholdings if desired",
        "Submit the birth certificate or adoption papers to HR",
        "Consider enrolling in or updating dependent care FSA",
        "Register for dependent care backup services",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _save(doc, "Leave-TimeOff", "LVE-001-Parental-Leave-Policy.docx")


def create_fmla_guide():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("LVE-002: FMLA Guide", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "LVE-002"),
        ("Owner", "VP of Total Rewards"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "The Family and Medical Leave Act (FMLA) provides eligible employees with up to 12 weeks of "
        "unpaid, job-protected leave per year for specified family and medical reasons."
    )

    doc.add_heading("2. Eligibility", level=2)
    for item in [
        "Employed for at least 12 months (not necessarily consecutive)",
        "Worked at least 1,250 hours in the 12 months prior to leave",
        "Work at a location where the company has at least 50 employees within 75 miles",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Qualifying Reasons", level=2)
    for item in [
        "Birth, adoption, or foster placement of a child",
        "Care for a spouse, child, or parent with a serious health condition",
        "Employee's own serious health condition that prevents performing job duties",
        "Qualifying exigency related to a family member's military service",
        "Care for a covered servicemember with a serious injury (up to 26 weeks)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Key Provisions", level=2)
    for item in [
        "Leave may be taken continuously, intermittently, or on a reduced schedule",
        "Health insurance continues at the same employee premium rate",
        "Job protection: return to the same or equivalent position",
        "30 days advance notice required when leave is foreseeable",
        "Medical certification required for health-condition-based leave",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _save(doc, "Leave-TimeOff", "LVE-002-FMLA-Guide.docx")


def create_pto_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("LVE-003: PTO & Bereavement Leave Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "LVE-003"),
        ("Owner", "VP of Total Rewards"),
    ])

    doc.add_heading("1. PTO Accrual Schedule", level=2)
    _add_table(doc,
        ["Years of Service", "Annual PTO Days", "Accrual Rate (per pay period)"],
        [
            ["0-2 years", "15 days", "4.62 hours"],
            ["3-5 years", "20 days", "6.15 hours"],
            ["6-10 years", "25 days", "7.69 hours"],
            ["10+ years", "30 days", "9.23 hours"],
        ]
    )

    doc.add_heading("2. PTO Policies", level=2)
    for item in [
        "Maximum carryover: 5 days into the next year",
        "PTO payout at termination: accrued, unused PTO is paid out in the final paycheck",
        "Minimum increment: PTO may be taken in 1-hour increments",
        "Advance notice: 2 weeks for planned absences of 3+ days",
        "Manager approval: required for all PTO requests",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Bereavement Leave", level=2)
    _add_table(doc,
        ["Relationship", "Paid Days"],
        [
            ["Spouse/domestic partner, child, parent", "5 days"],
            ["Sibling, grandparent, grandchild, in-law", "3 days"],
            ["Extended family, close friend", "1 day"],
        ]
    )
    doc.add_paragraph("Additional unpaid leave may be granted at manager's discretion. Bereavement leave does not reduce PTO balance.")

    _save(doc, "Leave-TimeOff", "LVE-003-PTO-Bereavement-Policy.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 5. PERSONAL DATA CHANGES
# ═══════════════════════════════════════════════════════════════════════════════

def create_name_change_procedure():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("PDC-001: Name Change Procedure", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "PDC-001"),
        ("Owner", "HR Operations Manager"),
        ("Risk Level", "HIGH – Requires documentation and HR Ops approval"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "Legal name changes affect payroll tax records, benefits enrollment, corporate identity systems (Active Directory, email), "
        "and compliance records. All legal name changes require verified documentation and dual approval."
    )

    doc.add_heading("2. Required Documentation", level=2)
    _add_table(doc,
        ["Reason for Change", "Accepted Documents"],
        [
            ["Marriage", "Certified marriage certificate"],
            ["Court order", "Court-ordered name change decree"],
            ["Divorce", "Divorce decree showing name reversion"],
            ["Personal preference (legal)", "Court-ordered name change decree"],
            ["Gender identity", "Court order or updated government-issued ID"],
        ]
    )

    doc.add_heading("3. Process", level=2)
    for i, step in enumerate([
        "Submit a name change request through the HR Portal or HR Concierge.",
        "Upload a scanned copy of the required documentation.",
        "HR Operations reviews the documentation (1-2 business days).",
        "If approved, changes are made in the following systems: HRIS/Workday, Payroll, Active Directory/Email, Benefits enrollment, Corporate directory, Building access/ID badge.",
        "Employee receives a confirmation email with a summary of all updated systems.",
        "New ID badge is issued within 5 business days.",
    ], 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("4. Processing Timeline", level=2)
    doc.add_paragraph(
        "Total processing time: 7-10 business days from submission of complete documentation. "
        "Systems are updated in the following order: HRIS (Day 1-2) → Payroll (Day 2-3) → IT/AD (Day 3-5) → "
        "Benefits (Day 5-7) → ID Badge (Day 7-10)."
    )

    doc.add_heading("5. Preferred Name vs. Legal Name", level=2)
    doc.add_paragraph(
        "If you want to use a different name in day-to-day work without changing your legal name, see the "
        "Preferred Name Update policy (PDC-005). Preferred name changes are self-service, require no documentation, "
        "and take effect within 24-48 hours."
    )

    _save(doc, "Personal-Data-Changes", "PDC-001-Name-Change-Procedure.docx")


def create_address_change_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("PDC-002: Address Change Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "PDC-002"),
        ("Owner", "HR Operations Manager"),
        ("Risk Level", "LOW – Self-service (unless cross-state)"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "Address changes are generally self-service and take effect within 2-3 business days. "
        "Cross-state moves may have payroll tax implications that require Payroll review."
    )

    doc.add_heading("2. How to Update Your Address", level=2)
    for i, step in enumerate([
        "Log in to the HR Portal and navigate to Personal Information > Address.",
        "Enter your new address including street, city, state, ZIP code, and country.",
        "Enter your move date (effective date).",
        "Submit the change.",
    ], 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("3. Cross-State Moves", level=2)
    doc.add_paragraph("If you are moving to a different state, additional steps apply:")
    for item in [
        "Payroll will review and update your state tax withholding",
        "If your new state has different tax rates, you may notice a change in net pay",
        "Remote work agreements may need to be updated",
        "Benefits availability may change based on your new location (HMO plans are location-dependent)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Processing Timeline", level=2)
    _add_table(doc,
        ["Change Type", "Timeline", "Approval"],
        [
            ["Same-state address change", "2-3 business days", "Self-service (no approval)"],
            ["Cross-state move", "5-7 business days", "Payroll review required"],
            ["International relocation", "Contact HR Global Mobility", "HR + Legal + Tax review"],
        ]
    )

    _save(doc, "Personal-Data-Changes", "PDC-002-Address-Change-Policy.docx")


def create_emergency_contact_guide():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("PDC-003: Emergency Contact Update Guide", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "PDC-003"),
        ("Owner", "HR Operations Manager"),
        ("Risk Level", "LOW – Self-service, immediate effect"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "Keep your emergency contact information up to date so we can reach someone on your behalf "
        "in an emergency. You may designate up to 3 emergency contacts."
    )

    doc.add_heading("2. Information Needed", level=2)
    for item in [
        "Contact full name",
        "Relationship to you (spouse, parent, sibling, friend, other)",
        "Phone number (at least one; mobile preferred)",
        "Email address (optional)",
        "Whether this is your primary contact",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. How to Update", level=2)
    doc.add_paragraph(
        "Update your emergency contacts through the HR Portal under Personal Information > Emergency Contacts. "
        "Changes take effect immediately. No approval is required."
    )

    _save(doc, "Personal-Data-Changes", "PDC-003-Emergency-Contact-Guide.docx")


def create_gov_id_update_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("PDC-004: Government ID & Passport Update Requirements", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "PDC-004"),
        ("Owner", "HR Operations Manager & Compliance"),
        ("Risk Level", "HIGH – Requires verified documents and dual approval"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "Government ID and passport information must be kept current for compliance, I-9 verification, "
        "and identity management purposes."
    )

    doc.add_heading("2. When to Update", level=2)
    for item in [
        "Passport renewal or new passport issued",
        "Name change on government-issued ID",
        "Driver's license renewal or state change",
        "Work authorization document renewal",
        "Permanent resident card renewal",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Required Documents", level=2)
    doc.add_paragraph(
        "You must provide a clear copy (scan or photo) of the new document. Original documents may be "
        "requested for in-person verification. See the I-9 Acceptable Documents list for details."
    )

    doc.add_heading("4. Processing", level=2)
    doc.add_paragraph(
        "Processing time: 7-10 business days. Government ID updates require dual approval from your manager "
        "and HR Operations. Compliance team may also review for I-9 purposes."
    )

    _save(doc, "Personal-Data-Changes", "PDC-004-Government-ID-Update-Requirements.docx")


def create_preferred_name_policy():
    doc = Document()
    _style_doc(doc)

    doc.add_heading("PDC-005: Preferred Name Update Policy", level=1)
    _add_metadata_table(doc, [
        ("Policy ID", "PDC-005"),
        ("Owner", "HR Operations Manager"),
        ("Risk Level", "LOW – Self-service, no documentation required"),
    ])

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "We respect every employee's right to be addressed by their chosen name. Preferred name changes "
        "update display systems only and do not affect legal or payroll records."
    )

    doc.add_heading("2. What Gets Updated", level=2)
    for item in [
        "Email display name",
        "Microsoft Teams profile",
        "Corporate directory / org chart",
        "Building access display (optional)",
        "Business cards (upon request)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. What Does NOT Change", level=2)
    for item in [
        "Legal name on payroll and tax documents",
        "Benefits enrollment records",
        "Government ID records",
        "Formal HR/legal correspondence",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. How to Update", level=2)
    doc.add_paragraph(
        "Self-service through the HR Portal under Personal Information > Preferred Name. "
        "No documentation or approval required. Changes take effect within 24-48 hours."
    )

    _save(doc, "Personal-Data-Changes", "PDC-005-Preferred-Name-Policy.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Generating SharePoint HR policy documents (.docx)...\n")

    # HR Policies
    print("HR-Policies/")
    create_grievance_policy()
    create_code_of_conduct()
    create_anti_discrimination_policy()
    create_harassment_policy()
    create_retaliation_policy()
    create_employee_relations_procedures()

    # Benefits & Insurance
    print("\nBenefits-Insurance/")
    create_health_insurance_guide()
    create_dental_vision_guide()
    create_life_insurance_policy()
    create_401k_guide()
    create_qualifying_life_events_guide()

    # Payroll & Tax
    print("\nPayroll-Tax/")
    create_direct_deposit_policy()
    create_tax_withholding_guide()
    create_payroll_change_sla()

    # Leave & Time Off
    print("\nLeave-TimeOff/")
    create_parental_leave_policy()
    create_fmla_guide()
    create_pto_policy()

    # Personal Data Changes
    print("\nPersonal-Data-Changes/")
    create_name_change_procedure()
    create_address_change_policy()
    create_emergency_contact_guide()
    create_gov_id_update_policy()
    create_preferred_name_policy()

    print(f"\n✅ Done! All documents generated in: {BASE}")
